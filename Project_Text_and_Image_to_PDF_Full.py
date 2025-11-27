import io
import os
import zipfile
from typing import List, Dict, Tuple

import streamlit as st
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit
from PIL import Image
from pdf2image import convert_from_bytes
from docx import Document
import csv

# ---------- Helpers ----------
A4_WIDTH, A4_HEIGHT = A4
MARGIN = 40  # left/right/top/bottom

def register_font_from_uploaded(uploaded_font) -> Tuple[str,bool]:
    if uploaded_font is None:
        return "Helvetica", False
    try:
        font_bytes = uploaded_font.read()
        font_filename = f"uploaded_{uploaded_font.name}"
        with open(font_filename,"wb") as f:
            f.write(font_bytes)
        font_name = os.path.splitext(uploaded_font.name)[0]
        pdfmetrics.registerFont(TTFont(font_name,font_filename))
        return font_name, True
    except:
        return "Helvetica", False

def color_tuple_from_hex(hex_color:str)->Tuple[float,float,float]:
    hex_color = hex_color.lstrip("#")
    r=int(hex_color[0:2],16)/255.0
    g=int(hex_color[2:4],16)/255.0
    b=int(hex_color[4:6],16)/255.0
    return (r,g,b)

def add_text_to_canvas(c:canvas.Canvas, text:str, font_name:str,font_size:int,color_rgb:Tuple[float,float,float]):
    c.setFont(font_name,font_size)
    c.setFillColor(colors.Color(*color_rgb))
    max_width = A4_WIDTH - 2*MARGIN
    paragraphs = text.splitlines()
    y = getattr(c,"_current_y", A4_HEIGHT-MARGIN)
    line_height = font_size*1.2
    for para in paragraphs:
        if para.strip()=="":
            y-=line_height
            if y<MARGIN:
                c.showPage()
                c.setFont(font_name,font_size)
                c.setFillColor(colors.Color(*color_rgb))
                y=A4_HEIGHT-MARGIN
            continue
        wrapped = simpleSplit(para,font_name,font_size,max_width)
        for wline in wrapped:
            if y<MARGIN:
                c.showPage()
                c.setFont(font_name,font_size)
                c.setFillColor(colors.Color(*color_rgb))
                y=A4_HEIGHT-MARGIN
            c.drawString(MARGIN,y,wline)
            y-=line_height
    c._current_y = y if y>MARGIN else (A4_HEIGHT-MARGIN)

def add_image_to_canvas(c:canvas.Canvas, image_bytes:bytes):
    c.showPage()
    img = Image.open(io.BytesIO(image_bytes))
    if img.mode not in ("RGB","RGBA"):
        img=img.convert("RGB")
    img_w,img_h = img.size
    max_w = A4_WIDTH-2*MARGIN
    max_h = A4_HEIGHT-2*MARGIN
    scale = min(max_w/img_w, max_h/img_h,1.0)
    draw_w = img_w*scale
    draw_h = img_h*scale
    x=(A4_WIDTH-draw_w)/2
    y=(A4_HEIGHT-draw_h)/2
    c.drawImage(ImageReader(img),x,y,width=draw_w,height=draw_h, preserveAspectRatio=True,mask='auto')
    c._current_y = A4_HEIGHT-MARGIN

def build_pdf(items:List[Dict], font_name:str,font_size:int,color_rgb:Tuple[float,float,float])->bytes:
    buffer=io.BytesIO()
    c=canvas.Canvas(buffer,pagesize=A4)
    c._current_y = A4_HEIGHT-MARGIN
    c.setFont(font_name,font_size)
    c.setFillColor(colors.Color(*color_rgb))
    for item in items:
        if item["type"]=="text":
            text=item["content"]
            if isinstance(text,bytes):
                try: text=text.decode("utf-8")
                except: text=text.decode("latin-1",errors="ignore")
            add_text_to_canvas(c,text,font_name,font_size,color_rgb)
        elif item["type"]=="image":
            add_image_to_canvas(c,item["content"])
    c.save()
    buffer.seek(0)
    return buffer.read()

def build_docx(items:List[Dict])->bytes:
    doc = Document()
    for item in items:
        if item["type"]=="text":
            text=item["content"]
            if isinstance(text,bytes):
                try: text=text.decode("utf-8")
                except: text=text.decode("latin-1",errors="ignore")
            doc.add_paragraph(text)
        elif item["type"]=="image":
            img = io.BytesIO(item["content"])
            doc.add_picture(img)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def pdf_to_images(pdf_bytes:bytes)->List[bytes]:
    try:
        images = convert_from_bytes(pdf_bytes)
        result=[]
        for im in images:
            bio=io.BytesIO()
            im.save(bio,format="JPEG")
            bio.seek(0)
            result.append(bio.read())
        return result
    except:
        return []

def build_csv(items:List[Dict])->bytes:
    bio=io.BytesIO()
    writer = csv.writer(io.TextIOWrapper(bio,"utf-8",newline=""))
    writer.writerow(["Type","Name","Content"])
    for it in items:
        content = it["content"]
        if isinstance(content,bytes): content = content.decode("utf-8","ignore") if it["type"]=="text" else "<IMAGE_BINARY>"
        writer.writerow([it["type"],it["name"],content])
    bio.seek(0)
    return bio.read()

# ---------- Streamlit UI ----------
st.set_page_config(page_title="Multi Text+Image → PDF/DOCX/CSV",layout="wide")
st.title("📄 Multi Text + Image → PDF/Word/CSV")

with st.expander("Upload files"):
    uploaded_texts = st.file_uploader("Upload .txt files", type=["txt"], accept_multiple_files=True)
    uploaded_images = st.file_uploader("Upload images", type=["jpg","jpeg","png","bmp"], accept_multiple_files=True)

col1,col2 = st.columns(2)
with col1:
    uploaded_font = st.file_uploader("Optional: upload .ttf font", type=["ttf"])
    font_name,font_uploaded = register_font_from_uploaded(uploaded_font)
with col2:
    font_size = st.number_input("Font size",6,48,12)
    color = st.color_picker("Text color","#000000")
    color_rgb = color_tuple_from_hex(color)

# ---------- Prepare items ----------
items=[]
if uploaded_texts:
    st.subheader("Edit Texts")
    for i,f in enumerate(uploaded_texts):
        content = f.read().decode("utf-8")
        new_text = st.text_area(f.name, content, key=f"text_{i}")
        items.append({"type":"text","name":f.name,"content":new_text})

if uploaded_images:
    st.subheader("Images")
    for i,f in enumerate(uploaded_images):
        st.image(f,caption=f.name,use_column_width=True)
        del_chk = st.checkbox(f"Delete {f.name}",key=f"del_img_{i}")
        if not del_chk:
            items.append({"type":"image","name":f.name,"content":f.read()})

# ---------- Reorder using experimental_data_editor ----------
if items:
    st.subheader("📌 Reorder files (drag & drop rows)")
    df = pd.DataFrame([{"Name":it["name"],"Type":it["type"]} for it in items])
    df_editor = st.experimental_data_editor(df,num_rows="dynamic",use_container_width=True)
    name_order = df_editor["Name"].tolist()
    items_ordered=[]
    for name in name_order:
        for it in items:
            if it["name"]==name:
                items_ordered.append(it)
                break
    items = items_ordered

# ---------- Build & Export ----------
if st.button("Build PDF & Preview"):
    pdf_bytes = build_pdf(items,font_name,font_size,color_rgb)
    st.success("✅ PDF built")
    # preview first page
    preview_imgs = pdf_to_images(pdf_bytes)
    if preview_imgs: st.image(preview_imgs[0], caption="Preview first page")
    st.download_button("⬇️ Download PDF",pdf_bytes,"merged.pdf","application/pdf")
    docx_bytes = build_docx(items)
    st.download_button("⬇️ Download Word",docx_bytes,"merged.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    csv_bytes = build_csv(items)
    st.download_button("⬇️ Download CSV",csv_bytes,"items.csv","text/csv")

if st.button("Export PDF pages to JPG (zip)"):
    if 'pdf_bytes' not in locals():
        st.warning("Please build PDF first")
    else:
        images = pdf_to_images(pdf_bytes)
        if images:
            zip_bio = io.BytesIO()
            with zipfile.ZipFile(zip_bio,"w") as zf:
                for idx,img_bytes in enumerate(images):
                    zf.writestr(f"page_{idx+1}.jpg",img_bytes)
            zip_bio.seek(0)
            st.download_button("⬇️ Download JPGs (zip)",zip_bio,"pages.zip","application/zip")
